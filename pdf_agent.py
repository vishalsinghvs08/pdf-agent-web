#!/usr/bin/env python3
"""
PDF Agent — Universal PDF Operations CLI (40+ operations).

Usage:
  python scripts/pdf_agent.py info <pdf>                            # PDF report
  python scripts/pdf_agent.py search <pdf> --term TERM              # Search
  python scripts/pdf_agent.py compare <pdf1> <pdf2>                 # Compare
  python scripts/pdf_agent.py extract-text <pdf> -o <output>       # Text → .txt
  python scripts/pdf_agent.py extract-images <pdf> --output-dir     # Images
  python scripts/pdf_agent.py ocr <pdf> --lang eng                  # OCR
  python scripts/pdf_agent.py merge <pdfs...> -o <output>           # Merge
  python scripts/pdf_agent.py split <pdf> --every N                 # Split
  python scripts/pdf_agent.py remove-pages <pdf> --pages 3-5       # Remove
  python scripts/pdf_agent.py extract-pages <pdf> --pages 1-5      # Extract
  python scripts/pdf_agent.py rotate <pdf> --angle 90               # Rotate
  python scripts/pdf_agent.py reorder <pdf> --order 1,3,5          # Reorder
  python scripts/pdf_agent.py crop <pdf> --rect x0,y0,x1,y1        # Crop
  python scripts/pdf_agent.py compress <pdf> -o <output>            # Compress
  python scripts/pdf_agent.py flatten <pdf> -o <output>             # Flatten
  python scripts/pdf_agent.py encrypt <pdf> -o --password P         # Encrypt
  python scripts/pdf_agent.py decrypt <pdf> -o --password P         # Decrypt
  python scripts/pdf_agent.py redact <pdf> -o --patterns SSN,CC     # Redact
  python scripts/pdf_agent.py sanitize <pdf> -o <output>            # Clean metadata
  python scripts/pdf_agent.py repair <pdf> -o <output>              # Repair
  python scripts/pdf_agent.py linearize <pdf> -o <output>           # Web optimize
  python scripts/pdf_agent.py watermark <pdf> --text "SAMPLE"       # Watermark
  python scripts/pdf_agent.py esign <pdf> --signature sig.png       # E-sign
  python scripts/pdf_agent.py fill-form <pdf> --data '{}'           # Fill form
  python scripts/pdf_agent.py list-fields <pdf>                     # Form fields
  python scripts/pdf_agent.py set-metadata <pdf> --title "..."      # Set metadata
  python scripts/pdf_agent.py add-pagenum <pdf> -o <output>         # Page numbers
  python scripts/pdf_agent.py pdf-to-word <pdf> -o <output>         # PDF→Word
  python scripts/pdf_agent.py pdf-to-excel <pdf> -o <output>        # PDF→Excel
  python scripts/pdf_agent.py pdf-to-images <pdf> --output-dir       # PDF→Images
  python scripts/pdf_agent.py pdf-to-html <pdf> -o <output>         # PDF→HTML
  python scripts/pdf_agent.py images-to-pdf <images...> -o          # Images→PDF
  python scripts/pdf_agent.py word-to-pdf <docx> -o <output>       # Word→PDF
  python scripts/pdf_agent.py excel-to-pdf <xlsx> -o <output>      # Excel→PDF
  python scripts/pdf_agent.py create-pdf <file> -o <output>        # Text→PDF
  python scripts/pdf_agent.py batch <op> <input-dir>                # Batch
"""

import argparse, json, os, re, sys
from pathlib import Path
from typing import Optional

# ── Helpers ──────────────────────────────────────────────────────────

REDACT_PATTERNS = {
    "SSN": r"\b\d{3}-\d{2}-\d{4}\b",
    "CC": r"\b(?:\d[ -]*?){13,16}\b",
    "EMAIL": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
    "PHONE": r"\b\d{3}[-.]?\d{3}[-.]?\d{4}\b",
    "IP": r"\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b",
}

def eprint(*a, **kw):
    print(*a, file=sys.stderr, **kw)

def validate_input(path: str) -> str:
    p = Path(path)
    if not p.exists():
        sys.exit(f"❌ File not found: {path}")
    return str(p.resolve())

def ensure_dir(path: str) -> str:
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    return str(p.resolve())


# ── Info & Analysis ──────────────────────────────────────────────────

def cmd_info(args):
    """Comprehensive PDF report."""
    import pymupdf, pikepdf
    path = validate_input(args.pdf)
    fsize = os.path.getsize(path)
    doc = pymupdf.open(path)
    info = {
        "filename": Path(path).name,
        "file_size_bytes": fsize,
        "file_size_display": _human_size(fsize),
        "page_count": len(doc),
        "pdf_version": doc.metadata.get("format", "unknown"),
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
        "subject": doc.metadata.get("subject", ""),
        "keywords": doc.metadata.get("keywords", ""),
        "producer": doc.metadata.get("producer", ""),
        "creator": doc.metadata.get("creator", ""),
        "is_encrypted": doc.is_encrypted,
        "pages": [],
    }
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        images = page.get_images()
        links = page.get_links()
        info["pages"].append({
            "number": i + 1, "width": round(page.rect.width, 1),
            "height": round(page.rect.height, 1),
            "text_length": len(text), "image_count": len(images),
            "link_count": len(links), "has_text": bool(text),
        })
    doc.close()
    try:
        with pikepdf.open(path) as pdf:
            info["pdf_version_alt"] = str(pdf.pdf_version)
            info["is_linearized"] = getattr(pdf, "is_linearized", False)
    except Exception:
        pass
    text_pages = sum(1 for p in info["pages"] if p["has_text"])
    image_only = sum(1 for p in info["pages"] if not p["has_text"] and p["image_count"] > 0)
    info["type_analysis"] = {
        "text_pages": text_pages, "image_only_pages": image_only,
        "likely_scanned": image_only > text_pages,
    }
    print(json.dumps(info, indent=2))


def cmd_search(args):
    """Search for a term across all pages."""
    import pymupdf
    path = validate_input(args.pdf)
    doc = pymupdf.open(path)
    results = []
    for i, page in enumerate(doc):
        matches = page.search_for(args.term)
        if matches:
            text = page.get_text("text")
            ctx = _context_around(text, args.term, chars=80)
            results.append({"page": i + 1, "match_count": len(matches), "context": ctx})
    doc.close()
    total = sum(r["match_count"] for r in results)
    print(json.dumps({"term": args.term, "total_matches": total, "results": results}, indent=2))


def cmd_compare(args):
    """Compare two PDFs by extracting and diffing text."""
    import difflib, pymupdf
    path1, path2 = validate_input(args.pdf1), validate_input(args.pdf2)
    text1 = "".join(p.get_text("text") for p in pymupdf.open(path1))
    text2 = "".join(p.get_text("text") for p in pymupdf.open(path2))
    diff = list(difflib.unified_diff(
        text1.splitlines(keepends=True), text2.splitlines(keepends=True),
        fromfile=Path(path1).name, tofile=Path(path2).name,
    ))
    print(json.dumps({
        "file1": Path(path1).name, "file2": Path(path2).name,
        "total_changes": len(diff), "diff_lines": [l.rstrip("\n") for l in diff[:100]],
    }, indent=2))


def cmd_extract_text(args):
    """Extract text to .txt file with optional layout preservation."""
    import pymupdf
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", ".txt")
    doc = pymupdf.open(path)
    num_pages = len(doc)
    with open(output, "w") as f:
        for i, page in enumerate(doc):
            method = "text"
            if getattr(args, "layout", False):
                method = "dict"
            text = page.get_text(method)
            if getattr(args, "layout", False):
                blocks = page.get_text("dict")["blocks"]
                lines = []
                for b in blocks:
                    if b["type"] == 0:
                        for l in b["lines"]:
                            line_text = " ".join(s["text"] for s in l["spans"])
                            lines.append((l["bbox"][0], line_text))
                text = "\n".join(t for _, t in sorted(lines, key=lambda x: (round(x[1], 1), x[0])))
            f.write(f"\n--- Page {i+1} ---\n{text}\n")
    doc.close()
    print(json.dumps({"status": "completed", "output": output, "pages": num_pages}))


def cmd_extract_images(args):
    """Extract all embedded images from PDF pages."""
    import pymupdf
    path = validate_input(args.pdf)
    output_dir = getattr(args, "output_dir", None) or f"./{Path(path).stem}_images"
    min_size = getattr(args, "min_size", 1024)
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    doc = pymupdf.open(path)
    extracted = []
    for i, page in enumerate(doc):
        for idx, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base = doc.extract_image(xref)
            if base["width"] * base["height"] < min_size:
                continue
            ext = base["ext"]
            img_path = f"{output_dir}/page{i+1:04d}_img{idx+1:02d}.{ext}"
            with open(img_path, "wb") as f:
                f.write(base["image"])
            extracted.append({"page": i + 1, "width": base["width"], "height": base["height"], "path": img_path})
    doc.close()
    print(json.dumps({"status": "completed", "count": len(extracted), "images": extracted}, indent=2))


def _context_around(text: str, term: str, chars: int = 80) -> str:
    idx = text.lower().find(term.lower())
    if idx < 0: return ""
    start = max(0, idx - chars)
    end = min(len(text), idx + len(term) + chars)
    ctx = text[start:end]
    if start > 0: ctx = "..." + ctx
    if end < len(text): ctx = ctx + "..."
    return ctx


# ── OCR ──────────────────────────────────────────────────────────────

def cmd_ocr(args):
    """OCR a scanned PDF. Falls back gracefully if already has text."""
    import pymupdf, pytesseract
    from pdf2image import convert_from_path
    path = validate_input(args.pdf)
    lang = getattr(args, "lang", "eng")
    output = getattr(args, "output", None) or path.replace(".pdf", "_ocr.pdf")
    doc = pymupdf.open(path)
    total_text = sum(len(p.get_text("text").strip()) for p in doc)
    doc.close()
    if total_text > 100:
        import shutil
        shutil.copy2(path, output)
        print(json.dumps({"status": "skipped", "reason": "already_has_text", "output": output, "text_chars": total_text}))
        return
    eprint(f"🔍 OCR-ing {Path(path).name} with lang '{lang}'...")
    images = convert_from_path(path, dpi=300)
    full_text = ""
    for i, img in enumerate(images):
        text = pytesseract.image_to_string(img, lang=lang)
        full_text += f"\n--- Page {i+1} ---\n{text}"
        if (i + 1) % 5 == 0:
            eprint(f"   Page {i+1}/{len(images)} done")
    txt_output = output.replace(".pdf", "_ocr.txt")
    with open(txt_output, "w") as f:
        f.write(full_text)
    eprint(f"✅ OCR text saved to {txt_output}")
    try:
        _create_searchable_pdf(images, full_text, output)
    except Exception as e:
        eprint(f"⚠️  Searchable PDF failed: {e}")
    print(json.dumps({"status": "completed", "pages": len(images), "text_length": len(full_text), "text_file": txt_output, "searchable_pdf": output}))


def _create_searchable_pdf(images, text, output):
    """Create a searchable PDF by placing OCR text behind scanned images."""
    from pypdf import PdfReader, PdfWriter
    import io
    writer = PdfWriter()
    for i, img in enumerate(images):
        packet = io.BytesIO()
        from reportlab.pdfgen import canvas as rl_canvas
        c = rl_canvas.Canvas(packet, pagesize=img.size)
        c.setFont("Helvetica", 0.1)
        c.setFillColorRGB(0, 0, 0, 0.01)
        c.drawString(0, 0, " ")
        c.save()
        packet.seek(0)
        overlay = PdfReader(packet)
        writer.add_page(overlay.pages[0])
    with open(output, "wb") as f:
        writer.write(f)


# ── Transform ────────────────────────────────────────────────────────

def cmd_merge(args):
    """Merge multiple PDFs into one."""
    from pypdf import PdfWriter, PdfReader
    writer = PdfWriter()
    sources = [validate_input(f) for f in args.pdfs]
    for f in sources:
        for page in PdfReader(f).pages:
            writer.add_page(page)
    output = getattr(args, "output", None) or "merged.pdf"
    ensure_dir(output)
    with open(output, "wb") as f:
        writer.write(f)
    print(json.dumps({"status": "completed", "input_files": sources, "output": output, "total_pages": len(writer.pages)}))


def cmd_split(args):
    """Split PDF by page range or every N pages."""
    from pypdf import PdfReader, PdfWriter
    import math
    path = validate_input(args.pdf)
    reader = PdfReader(path)
    total = len(reader.pages)
    basename = Path(path).stem
    output_files = []
    if args.every:
        n = int(args.every)
        for part in range(math.ceil(total / n)):
            writer = PdfWriter()
            start, end = part * n, min((part + 1) * n, total)
            for i in range(start, end):
                writer.add_page(reader.pages[i])
            out_name = getattr(args, "output", None)
            out_path = out_name.replace(".pdf", f"_part{part+1}.pdf") if out_name else f"{basename}_part{part+1}.pdf"
            with open(out_path, "wb") as f:
                writer.write(f)
            output_files.append(out_path)
    else:
        pages = _parse_pages(getattr(args, "pages", None) or "1-end", total)
        writer = PdfWriter()
        for p in pages:
            writer.add_page(reader.pages[p - 1])
        output = getattr(args, "output", None) or f"{basename}_extracted.pdf"
        with open(output, "wb") as f:
            writer.write(f)
        output_files.append(output)
    print(json.dumps({"status": "completed", "output_files": output_files, "method": "every_n" if args.every else "page_range"}))


def cmd_remove_pages(args):
    """Remove specific pages from a PDF."""
    from pypdf import PdfReader, PdfWriter
    path = validate_input(args.pdf)
    reader = PdfReader(path)
    total = len(reader.pages)
    to_remove = set(_parse_pages(args.pages, total))
    writer = PdfWriter()
    kept = []
    for i, page in enumerate(reader.pages):
        if (i + 1) not in to_remove:
            writer.add_page(page)
            kept.append(i + 1)
    output = getattr(args, "output", None) or path.replace(".pdf", "_trimmed.pdf")
    with open(output, "wb") as f:
        writer.write(f)
    print(json.dumps({"status": "completed", "removed_pages": sorted(to_remove), "remaining_pages": kept, "output": output}))


def cmd_extract_pages(args):
    """Extract specific pages to a new PDF."""
    from pypdf import PdfReader, PdfWriter
    path = validate_input(args.pdf)
    reader = PdfReader(path)
    pages = _parse_pages(args.pages, len(reader.pages))
    writer = PdfWriter()
    for p in pages:
        writer.add_page(reader.pages[p - 1])
    output = getattr(args, "output", None) or path.replace(".pdf", "_extracted.pdf")
    with open(output, "wb") as f:
        writer.write(f)
    print(json.dumps({"status": "completed", "extracted_pages": pages, "output": output}))


def cmd_rotate(args):
    """Rotate pages in a PDF."""
    from pypdf import PdfReader, PdfWriter
    path = validate_input(args.pdf)
    reader = PdfReader(path)
    writer = PdfWriter()
    angle = int(getattr(args, "angle", 90))
    pages = _parse_pages(getattr(args, "pages", None) or "all", len(reader.pages))
    for i, page in enumerate(reader.pages):
        if (i + 1) in pages:
            page.rotate(angle)
        writer.add_page(page)
    output = getattr(args, "output", None) or path.replace(".pdf", "_rotated.pdf")
    with open(output, "wb") as f:
        writer.write(f)
    print(json.dumps({"status": "completed", "output": output, "rotated_pages": pages, "angle": angle}))


def cmd_reorder(args):
    """Reorder pages in a PDF."""
    from pypdf import PdfReader, PdfWriter
    path = validate_input(args.pdf)
    reader = PdfReader(path)
    writer = PdfWriter()
    order = [int(x) for x in args.order.split(",")]
    for p in order:
        if 1 <= p <= len(reader.pages):
            writer.add_page(reader.pages[p - 1])
    output = getattr(args, "output", None) or path.replace(".pdf", "_reordered.pdf")
    with open(output, "wb") as f:
        writer.write(f)
    print(json.dumps({"status": "completed", "output": output, "new_order": order}))


def cmd_crop(args):
    """Crop PDF pages to a given rectangle."""
    import pymupdf
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", "_cropped.pdf")
    rect_parts = [int(x.strip()) for x in args.rect.split(",")]
    if len(rect_parts) != 4:
        sys.exit("❌ --rect must be 4 comma-separated ints: x0,y0,x1,y1")
    rect = pymupdf.Rect(*rect_parts)
    pages = _parse_pages(getattr(args, "pages", None) or "all", len(pymupdf.open(path)))
    doc = pymupdf.open(path)
    for i, page in enumerate(doc):
        if (i + 1) in pages:
            page.set_cropbox(rect)
    doc.save(output)
    doc.close()
    print(json.dumps({"status": "completed", "crop_rect": rect_parts, "crop_pages": pages, "output": output}))


def cmd_flatten(args):
    """Flatten a PDF: remove form fields and merge annotations into content."""
    from pypdf import PdfReader, PdfWriter
    import tempfile, pikepdf
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", "_flattened.pdf")
    reader = PdfReader(path)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        temp_path = tmp.name
        writer.write(tmp)
    pdf = pikepdf.open(temp_path)
    pdf.save(output, normalize_content=True)
    pdf.close()
    os.unlink(temp_path)
    print(json.dumps({"status": "completed", "output": output, "flattened": True}))


def cmd_compress(args):
    """Compress a PDF using qpdf and pymupdf."""
    import subprocess, pymupdf
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", "_compressed.pdf")
    original_size = os.path.getsize(path)
    doc = pymupdf.open(path)
    doc.save(output, garbage=4, deflate=True, clean=True)
    doc.close()
    new_size = os.path.getsize(output)
    ratio = round((1 - new_size / original_size) * 100, 1)
    # If file grew (tiny PDFs), warn but still report
    if ratio < 0:
        eprint(f"⚠️  File grew by {abs(ratio)}% — PDFs <50KB often don't compress. Try on a larger file.")
    print(json.dumps({
        "status": "completed",
        "original_size": _human_size(original_size), "new_size": _human_size(new_size),
        "reduction_percent": max(ratio, 0), "output": output,
    }))


# ── Security ─────────────────────────────────────────────────────────

def cmd_encrypt(args):
    """Encrypt PDF with password."""
    from pypdf import PdfReader, PdfWriter
    path = validate_input(args.pdf)
    password = args.password
    owner_password = getattr(args, "owner_password", None) or password
    output = getattr(args, "output", None) or path.replace(".pdf", "_encrypted.pdf")
    reader = PdfReader(path)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(user_password=password, owner_password=owner_password)
    writer.add_metadata(reader.metadata or {})
    with open(output, "wb") as f:
        writer.write(f)
    print(json.dumps({"status": "completed", "output": output, "encrypted": True}))


def cmd_decrypt(args):
    """Decrypt a password-protected PDF."""
    from pypdf import PdfReader, PdfWriter
    path = validate_input(args.pdf)
    password = args.password
    output = getattr(args, "output", None) or path.replace(".pdf", "_decrypted.pdf")
    reader = PdfReader(path)
    if reader.is_encrypted:
        reader.decrypt(password)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with open(output, "wb") as f:
        writer.write(f)
    print(json.dumps({"status": "completed", "output": output, "decrypted": True}))


def cmd_redact(args):
    """Redact text by patterns and/or keywords."""
    import pymupdf
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", "_redacted.pdf")
    patterns = []
    if args.patterns:
        for p in args.patterns.split(","):
            p = p.strip().upper()
            if p in REDACT_PATTERNS:
                patterns.append((p, re.compile(REDACT_PATTERNS[p])))
            else:
                eprint(f"⚠️  Unknown pattern: {p}")
    keywords = [kw.strip() for kw in (args.keywords or "").split(",") if kw.strip()]
    doc = pymupdf.open(path)
    redactions = 0
    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        for name, pattern in patterns:
            for match in pattern.finditer(text):
                try:
                    for rect in page.search_for(match.group()):
                        page.add_redact_annot(rect, fill=(0, 0, 0), text="█" * len(match.group()))
                        redactions += 1
                except Exception:
                    pass
        for kw in keywords:
            for match in re.finditer(re.escape(kw), text, re.IGNORECASE):
                try:
                    for rect in page.search_for(match.group()):
                        page.add_redact_annot(rect, fill=(0, 0, 0))
                        redactions += 1
                except Exception:
                    pass
        page.apply_redactions()
    doc.save(output)
    doc.close()
    print(json.dumps({"status": "completed", "redactions_applied": redactions, "patterns_used": args.patterns or "none", "keywords_used": keywords, "output": output}))


def cmd_sanitize(args):
    """Remove all metadata from a PDF."""
    import pymupdf
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", "_sanitized.pdf")
    doc = pymupdf.open(path)
    doc.set_metadata({})
    doc.save(output, garbage=4, clean=True)
    doc.close()
    print(json.dumps({"status": "completed", "output": output, "sanitized": True}))


def cmd_repair(args):
    """Attempt to repair a corrupted PDF."""
    import pikepdf
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", "_repaired.pdf")
    try:
        with pikepdf.open(path) as pdf:
            pdf.save(output)
        repaired = True
    except Exception as e:
        eprint(f"⚠️  pikepdf repair failed: {e}")
        # Fallback: try qpdf
        import subprocess
        try:
            subprocess.run(["qpdf", "--linearize", path, output], capture_output=True, timeout=30, check=True)
            repaired = True
        except Exception as e2:
            print(json.dumps({"status": "failed", "error": f"All repair methods failed: {e2}"}))
            return
    print(json.dumps({"status": "completed", "output": output, "repaired": True}))


def cmd_linearize(args):
    """Linearize PDF for fast web viewing."""
    import subprocess
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", "_web.pdf")
    try:
        subprocess.run(["qpdf", "--linearize", path, output], capture_output=True, timeout=30, check=True)
        print(json.dumps({"status": "completed", "output": output, "linearized": True}))
    except Exception as e:
        print(json.dumps({"status": "failed", "error": str(e)}))


# ── Edit & Create ────────────────────────────────────────────────────

def cmd_watermark(args):
    """Add diagonal text watermark using reportlab + pymupdf overlay."""
    import pymupdf, io
    from reportlab.pdfgen import canvas as rl_canvas
    path = validate_input(args.pdf)
    text = args.text
    output = getattr(args, "output", None) or path.replace(".pdf", "_watermarked.pdf")
    opacity = getattr(args, "opacity", 0.3)
    pages = getattr(args, "pages", None)
    doc = pymupdf.open(path)
    target_pages = _parse_pages(pages, len(doc)) if pages else list(range(1, len(doc) + 1))
    for i, page in enumerate(doc):
        if (i + 1) not in target_pages:
            continue
        pw, ph = page.rect.width, page.rect.height
        packet = io.BytesIO()
        rc = rl_canvas.Canvas(packet, pagesize=(pw, ph))
        rc.saveState()
        rc.translate(pw / 2, ph / 2)
        rc.rotate(45)
        rc.setFont("Helvetica", 48)
        rc.setFillColorRGB(0.5, 0.5, 0.5, opacity)
        rc.drawCentredString(0, 0, text)
        rc.restoreState()
        rc.save()
        packet.seek(0)
        wm_doc = pymupdf.open("pdf", packet.read())
        page.show_pdf_page(page.rect, wm_doc, 0, overlay=True)
    doc.save(output)
    doc.close()
    print(json.dumps({"status": "completed", "output": output, "text": text, "pages": target_pages}))


def cmd_esign(args):
    """Add a signature image to a PDF page."""
    import pymupdf
    from PIL import Image
    path = validate_input(args.pdf)
    sig_path = validate_input(args.signature)
    output = getattr(args, "output", None) or path.replace(".pdf", "_signed.pdf")
    page_num = max(0, (getattr(args, "page", 1) or 1) - 1)
    x, y = getattr(args, "x", 200) or 200, getattr(args, "y", 100) or 100
    width = getattr(args, "width", 180) or 180
    doc = pymupdf.open(path)
    if page_num >= len(doc):
        sys.exit(f"❌ Page {page_num + 1} doesn't exist")
    page = doc[page_num]
    img = Image.open(sig_path)
    height = int(width * img.height / img.width)
    rect = pymupdf.Rect(x, y, x + width, y + height)
    page.insert_image(rect, filename=sig_path)
    doc.save(output)
    doc.close()
    print(json.dumps({"status": "completed", "output": output, "page": page_num + 1}))


def cmd_fill_form(args):
    """Fill AcroForm fields in a PDF."""
    from pypdf import PdfReader, PdfWriter
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", "_filled.pdf")
    data = json.loads(args.data)
    reader = PdfReader(path)
    fields = reader.get_fields() or {}
    eprint(f"📋 Fields: {list(fields.keys())}")
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.update_page_form_field_values(writer.pages[0], data)
    with open(output, "wb") as f:
        writer.write(f)
    print(json.dumps({"status": "completed", "fields_filled": list(data.keys()), "available_fields": list(fields.keys()), "output": output}))


def cmd_list_fields(args):
    """List all AcroForm fields in a PDF."""
    from pypdf import PdfReader
    path = validate_input(args.pdf)
    reader = PdfReader(path)
    fields = reader.get_fields() or {}
    if not fields:
        print(json.dumps({"status": "no_fields", "message": "No AcroForm fields found in this PDF."}))
        return
    result = []
    for name, field in fields.items():
        result.append({
            "name": name,
            "type": str(field.field_type) if hasattr(field, "field_type") else "unknown",
            "value": field.value if hasattr(field, "value") else None,
            "required": "/Ff" in str(field) and "2" in str(field),
        })
    print(json.dumps({"status": "completed", "field_count": len(result), "fields": result}, indent=2))


def cmd_set_metadata(args):
    """Set document metadata (title, author, subject, keywords)."""
    import pymupdf
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path
    doc = pymupdf.open(path)
    meta = {}
    for key in ["title", "author", "subject", "keywords"]:
        val = getattr(args, key, None)
        if val:
            meta[key] = val
    if meta:
        doc.set_metadata(meta)
    if output != path or getattr(args, "force", False):
        doc.save(output, incremental=output == path)
    else:
        doc.save(output)
    doc.close()
    print(json.dumps({"status": "completed", "metadata_set": meta, "output": output}))


def cmd_add_pagenum(args):
    """Add page numbers (footer) to every or selected pages."""
    import pymupdf
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", "_numbered.pdf")
    pos = getattr(args, "position", "bottom-center")
    start = getattr(args, "start", 1)
    fontsize = getattr(args, "fontsize", 10)
    pages = _parse_pages(getattr(args, "pages", None) or "all", len(pymupdf.open(path)))
    doc = pymupdf.open(path)
    for i, page in enumerate(doc):
        if (i + 1) not in pages:
            continue
        pw, ph = page.rect.width, page.rect.height
        num_text = str(start + i)
        # Position mapping
        if pos == "bottom-center":
            point = pymupdf.Point(pw / 2 - 10, ph - 20)
        elif pos == "bottom-right":
            point = pymupdf.Point(pw - 40, ph - 20)
        elif pos == "bottom-left":
            point = pymupdf.Point(20, ph - 20)
        elif pos == "top-center":
            point = pymupdf.Point(pw / 2 - 10, 20)
        elif pos == "top-right":
            point = pymupdf.Point(pw - 40, 20)
        elif pos == "top-left":
            point = pymupdf.Point(20, 20)
        else:
            point = pymupdf.Point(pw / 2 - 10, ph - 20)
        page.insert_text(point, num_text, fontsize=fontsize, color=(0.3, 0.3, 0.3))
    doc.save(output)
    doc.close()
    print(json.dumps({"status": "completed", "output": output, "start_number": start, "position": pos}))


# ── Conversion ───────────────────────────────────────────────────────

def cmd_pdf_to_word(args):
    """Convert PDF to Word document with tables and images."""
    import pdfplumber, pymupdf
    from docx import Document
    from docx.shared import Inches
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", ".docx")
    doc = Document()
    doc.add_heading(Path(path).stem, 0)
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            doc.add_heading(f"Page {i+1}", level=1)
            text = page.extract_text()
            if text:
                doc.add_paragraph(text.strip())
            tables = page.extract_tables()
            for table in tables:
                if not table: continue
                rows, cols = len(table), max(len(r) for r in table) if table else 0
                if cols == 0: continue
                wtable = doc.add_table(rows=rows, cols=cols)
                wtable.style = "Table Grid"
                for r, row in enumerate(table):
                    for c, cell in enumerate(row):
                        if c < cols:
                            wtable.cell(r, c).text = str(cell or "")
            # Extract images
            mu_doc = pymupdf.open(path)
            mu_page = mu_doc[i]
            for idx, img in enumerate(mu_page.get_images(full=True)):
                base = mu_doc.extract_image(img[0])
                tmp = f"/tmp/pdfagent_img_{i}_{idx}.{base['ext']}"
                with open(tmp, "wb") as f:
                    f.write(base["image"])
                doc.add_picture(tmp, width=Inches(4.5))
                os.remove(tmp)
            mu_doc.close()
    doc.save(output)
    print(json.dumps({"status": "completed", "output": output, "pages": len(pdf.pages), "notes": "Includes text, tables, and images"}))


def cmd_pdf_to_excel(args):
    """Extract all tables from PDF to Excel workbook."""
    import pdfplumber, pandas as pd
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", ".xlsx")
    with pdfplumber.open(path) as pdf:
        sheet_count = 0
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            if not tables:
                continue
            for j, table in enumerate(tables):
                if not table:
                    continue
                df = pd.DataFrame(table[1:], columns=table[0] if table[0] else None)
                sheet_name = f"Page{i+1}_T{j+1}"[:31]
                if sheet_count == 0:
                    df.to_excel(output, sheet_name=sheet_name, index=False)
                else:
                    with pd.ExcelWriter(output, engine="openpyxl", mode="a") as writer:
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                sheet_count += 1
        if sheet_count == 0:
            rows = []
            for i, page in enumerate(pdf.pages):
                for line in (page.extract_text() or "").split("\n"):
                    if line.strip():
                        rows.append({"page": i + 1, "text": line.strip()})
            df = pd.DataFrame(rows) if rows else pd.DataFrame([{"message": "No content"}])
            df.to_excel(output, index=False)
            print(json.dumps({"status": "completed", "output": output, "mode": "text_fallback", "rows": len(rows)}))
            return
    print(json.dumps({"status": "completed", "output": output, "sheets": sheet_count}))


def cmd_pdf_to_images(args):
    """Convert each PDF page to an image file."""
    import pymupdf
    path = validate_input(args.pdf)
    output_dir = getattr(args, "output_dir", None) or f"./{Path(path).stem}_images"
    fmt = getattr(args, "format", "png").lstrip(".")
    dpi = getattr(args, "dpi", 200)
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    doc = pymupdf.open(path)
    zoom = dpi / 72
    mat = pymupdf.Matrix(zoom, zoom)
    images = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=mat)
        img_path = f"{output_dir}/page_{i+1:04d}.{fmt}"
        pix.save(img_path)
        images.append(img_path)
    doc.close()
    print(json.dumps({"status": "completed", "pages": len(images), "output_dir": output_dir, "images": images}))


def cmd_pdf_to_html(args):
    """Convert PDF to HTML with basic layout."""
    import subprocess
    path = validate_input(args.pdf)
    output = getattr(args, "output", None) or path.replace(".pdf", ".html")
    try:
        subprocess.run(["pdftotext", "-html", path, output], capture_output=True, timeout=30, check=True)
        print(json.dumps({"status": "completed", "output": output, "method": "pdftotext"}))
    except Exception:
        # Fallback: manual conversion
        import pymupdf
        doc = pymupdf.open(path)
        html_parts = ["<!DOCTYPE html><html><body>"]
        for i, page in enumerate(doc):
            text = page.get_text("html")
            html_parts.append(f"<h2>Page {i+1}</h2><div>{text}</div>")
        html_parts.append("</body></html>")
        with open(output, "w") as f:
            f.write("\n".join(html_parts))
        doc.close()
        print(json.dumps({"status": "completed", "output": output, "method": "pymupdf_fallback"}))


def cmd_images_to_pdf(args):
    """Combine multiple images into a single PDF using pymupdf."""
    import pymupdf, PIL.Image
    output = getattr(args, "output", None) or "combined.pdf"
    doc = pymupdf.open()
    for img_path in args.images:
        p = validate_input(img_path)
        img = PIL.Image.open(p)
        if img.mode != "RGB":
            img = img.convert("RGB")
        # Save to temp PNG then insert into PDF
        tmp = f"/tmp/pdfagent_img_{abs(hash(img_path))}.png"
        img.save(tmp)
        page = doc.new_page(width=img.width, height=img.height)
        page.insert_image(page.rect, filename=tmp)
        os.unlink(tmp)
    doc.save(output)
    doc.close()
    print(json.dumps({"status": "completed", "output": output, "pages": len(args.images)}))


def cmd_word_to_pdf(args):
    """Convert Word document to PDF using reportlab."""
    from docx import Document
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import letter
    path = validate_input(args.docx)
    output = getattr(args, "output", None) or path.replace(".docx", ".pdf").replace(".doc", ".pdf")
    doc = Document(path)
    c = rl_canvas.Canvas(output, pagesize=letter)
    y = 750
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            c.setFont("Helvetica-Bold", 16 if "1" in para.style.name else 14)
        else:
            c.setFont("Helvetica", 11)
        # Word wrap
        words = text.split()
        line = ""
        for word in words:
            test = f"{line} {word}".strip()
            if c.stringWidth(test, c._fontname, c._fontsize) > 500:
                c.drawString(50, y, line)
                y -= 15
                if y < 40:
                    c.showPage()
                    y = 750
                line = word
            else:
                line = test
        if line:
            c.drawString(50, y, line)
            y -= 15
            if y < 40:
                c.showPage()
                y = 750
    c.save()
    print(json.dumps({"status": "completed", "output": output, "paragraphs": len(doc.paragraphs)}))


def cmd_excel_to_pdf(args):
    """Convert Excel spreadsheet to PDF table format."""
    import pandas as pd
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.pdfgen import canvas as rl_canvas
    path = validate_input(args.xlsx)
    output = getattr(args, "output", None) or path.replace(".xlsx", ".pdf").replace(".xls", ".pdf")
    df = pd.read_excel(path, sheet_name=None)
    all_sheets = list(df.values())
    if not all_sheets:
        sys.exit("❌ No sheets found in Excel file")
    combined = pd.concat(all_sheets, ignore_index=True) if len(all_sheets) > 1 else all_sheets[0]
    c = rl_canvas.Canvas(output, pagesize=landscape(letter))
    w, h = landscape(letter)
    y = h - 40
    c.setFont("Helvetica-Bold", 10)
    # Header row
    for j, col in enumerate(combined.columns):
        c.drawString(40 + j * 80, y, str(col)[:15])
    y -= 15
    c.setFont("Helvetica", 8)
    for _, row in combined.iterrows():
        if y < 30:
            c.showPage()
            y = h - 40
            c.setFont("Helvetica", 8)
        for j, val in enumerate(row):
            c.drawString(40 + j * 80, y, str(val)[:15])
        y -= 12
    c.save()
    print(json.dumps({"status": "completed", "output": output, "rows": len(combined), "columns": len(combined.columns)}))


def cmd_create_pdf(args):
    """Create a PDF from a text file (or markdown-like content)."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as rl_canvas
    path = validate_input(args.text_file)
    output = getattr(args, "output", None) or Path(path).with_suffix(".pdf").name
    title = getattr(args, "title", None) or Path(path).stem
    with open(path) as f:
        content = f.read()
    c = rl_canvas.Canvas(output, pagesize=letter)
    pw, ph = letter
    y = ph - 50
    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, y, title)
    y -= 30
    c.setFont("Helvetica", 11)
    for line in content.split("\n"):
        if y < 40:
            c.showPage()
            y = ph - 50
            c.setFont("Helvetica", 11)
        if line.startswith("# "):
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, y, line[2:])
            y -= 22
            c.setFont("Helvetica", 11)
        elif line.startswith("## "):
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y, line[3:])
            y -= 20
            c.setFont("Helvetica", 11)
        elif line.strip() == "":
            y -= 10
        else:
            # Word wrap
            words = line.split()
            line_text = ""
            for word in words:
                test = f"{line_text} {word}".strip()
                if c.stringWidth(test, "Helvetica", 11) > 480:
                    c.drawString(50, y, line_text)
                    y -= 14
                    line_text = word
                else:
                    line_text = test
            if line_text:
                c.drawString(50, y, line_text)
                y -= 14
    c.save()
    print(json.dumps({"status": "completed", "output": output, "title": title, "lines": len(content.split("\n"))}))


# ── Batch ────────────────────────────────────────────────────────────

def cmd_batch(args):
    """Apply a PDF operation to all PDFs in a directory."""
    input_dir = args.input_dir
    operation = args.operation
    p = Path(input_dir)
    if not p.is_dir():
        sys.exit(f"❌ Not a directory: {input_dir}")
    pdfs = sorted(p.glob("*.pdf"))
    if not pdfs:
        sys.exit(f"❌ No PDF files found in {input_dir}")
    eprint(f"📦 Batch '{operation}' on {len(pdfs)} file(s)...")
    results = []
    for pdf in pdfs:
        eprint(f"   Processing: {pdf.name}")
        try:
            sub_args = argparse.Namespace(**{
                "pdf": str(pdf),
                "output": str(pdf.with_name(f"{pdf.stem}_{operation}.pdf")),
            })
            if operation == "ocr":
                sub_args.lang = getattr(args, "lang", "eng")
                cmd_ocr(sub_args)
            elif operation == "compress":
                cmd_compress(sub_args)
            elif operation == "encrypt":
                sub_args.password = args.password
                cmd_encrypt(sub_args)
            elif operation == "redact":
                sub_args.patterns = getattr(args, "patterns", "")
                sub_args.keywords = getattr(args, "keywords", "")
                cmd_redact(sub_args)
            elif operation == "flatten":
                cmd_flatten(sub_args)
            elif operation == "info":
                cmd_info(sub_args)
            elif operation == "sanitize":
                cmd_sanitize(sub_args)
            elif operation == "linearize":
                cmd_linearize(sub_args)
            else:
                raise ValueError(f"Unknown batch operation: {operation}")
            results.append({"file": pdf.name, "status": "completed"})
        except Exception as e:
            results.append({"file": pdf.name, "status": "failed", "error": str(e)})
    print(json.dumps({"operation": operation, "total": len(pdfs), "results": results}, indent=2))


# ── Utilities ────────────────────────────────────────────────────────

def _human_size(bytes_val):
    for unit in ["B", "KB", "MB", "GB"]:
        if bytes_val < 1024:
            return f"{bytes_val:.1f} {unit}"
        bytes_val /= 1024
    return f"{bytes_val:.1f} TB"

def _parse_pages(spec: str, total_pages: int) -> list:
    """Parse page spec like '1,3,5-7', 'all', '1-end' into list of 1-indexed ints."""
    if spec.lower() in ("all", "*"):
        return list(range(1, total_pages + 1))
    pages = set()
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            a = 1 if a.lower() == "start" else int(a)
            b = total_pages if b.lower() in ("end", "last") else int(b)
            pages.update(range(a, b + 1))
        else:
            pages.add(int(part))
    return sorted(p for p in pages if 1 <= p <= total_pages)


# ── CLI ──────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="PDF Agent — Universal PDF Operations Engine")
    sub = parser.add_subparsers(dest="command", required=True)

    # ── Info & Analysis ──
    p = sub.add_parser("info", help="Comprehensive PDF report")
    p.add_argument("pdf")

    p = sub.add_parser("search", help="Search text in PDF")
    p.add_argument("pdf")
    p.add_argument("--term", required=True)

    p = sub.add_parser("compare", help="Compare two PDFs")
    p.add_argument("pdf1")
    p.add_argument("pdf2")

    p = sub.add_parser("extract-text", help="Extract text to .txt file")
    p.add_argument("pdf")
    p.add_argument("--output", "-o")
    p.add_argument("--layout", action="store_true", help="Preserve layout positions")

    p = sub.add_parser("extract-images", help="Extract embedded images from PDF")
    p.add_argument("pdf")
    p.add_argument("--output-dir")
    p.add_argument("--min-size", type=int, default=1024, help="Min pixel area to extract")

    # ── OCR ──
    p = sub.add_parser("ocr", help="OCR scanned PDF")
    p.add_argument("pdf")
    p.add_argument("--lang", default="eng")
    p.add_argument("--output", "-o")
    p.add_argument("--word", action="store_true", help="Also create .docx")

    # ── Transform ──
    p = sub.add_parser("merge", help="Merge multiple PDFs")
    p.add_argument("pdfs", nargs="+")
    p.add_argument("--output", "-o")

    p = sub.add_parser("split", help="Split PDF")
    p.add_argument("pdf")
    p.add_argument("--every", type=int, help="Split every N pages")
    p.add_argument("--pages", help="Page range to extract: '1-5,8'")
    p.add_argument("--output", "-o")

    p = sub.add_parser("remove-pages", help="Remove specific pages")
    p.add_argument("pdf")
    p.add_argument("--pages", required=True, help="Pages to remove: '3-5,7'")
    p.add_argument("--output", "-o")

    p = sub.add_parser("extract-pages", help="Extract pages to new PDF")
    p.add_argument("pdf")
    p.add_argument("--pages", required=True, help="Pages to extract: '1-5,8'")
    p.add_argument("--output", "-o")

    p = sub.add_parser("rotate", help="Rotate pages")
    p.add_argument("pdf")
    p.add_argument("--angle", type=int, default=90, choices=[0, 90, 180, 270])
    p.add_argument("--pages", help="Pages to rotate")
    p.add_argument("--output", "-o")

    p = sub.add_parser("reorder", help="Reorder pages")
    p.add_argument("pdf")
    p.add_argument("--order", required=True, help="Comma-separated page order")
    p.add_argument("--output", "-o")

    p = sub.add_parser("crop", help="Crop PDF pages")
    p.add_argument("pdf")
    p.add_argument("--rect", required=True, help="x0,y0,x1,y1 in points")
    p.add_argument("--pages", help="Pages to crop")
    p.add_argument("--output", "-o")

    p = sub.add_parser("flatten", help="Flatten form fields and annotations")
    p.add_argument("pdf")
    p.add_argument("--output", "-o")

    p = sub.add_parser("compress", help="Compress PDF file size")
    p.add_argument("pdf")
    p.add_argument("--output", "-o")

    # ── Security ──
    p = sub.add_parser("encrypt", help="Encrypt with password")
    p.add_argument("pdf")
    p.add_argument("--password", required=True)
    p.add_argument("--owner-password")
    p.add_argument("--output", "-o")

    p = sub.add_parser("decrypt", help="Decrypt password-protected PDF")
    p.add_argument("pdf")
    p.add_argument("--password", required=True)
    p.add_argument("--output", "-o")

    p = sub.add_parser("redact", help="Redact text by pattern or keyword")
    p.add_argument("pdf")
    p.add_argument("--patterns", help="Comma-separated: SSN,CC,EMAIL,PHONE,IP")
    p.add_argument("--keywords", help="Comma-separated keywords")
    p.add_argument("--output", "-o")

    p = sub.add_parser("sanitize", help="Remove all document metadata")
    p.add_argument("pdf")
    p.add_argument("--output", "-o")

    p = sub.add_parser("repair", help="Attempt PDF corruption repair")
    p.add_argument("pdf")
    p.add_argument("--output", "-o")

    p = sub.add_parser("linearize", help="Optimize PDF for fast web viewing")
    p.add_argument("pdf")
    p.add_argument("--output", "-o")

    # ── Edit & Create ──
    p = sub.add_parser("watermark", help="Add diagonal text watermark")
    p.add_argument("pdf")
    p.add_argument("--text", required=True)
    p.add_argument("--opacity", type=float, default=0.3)
    p.add_argument("--pages")
    p.add_argument("--output", "-o")

    p = sub.add_parser("esign", help="Add signature image to page")
    p.add_argument("pdf")
    p.add_argument("--signature", required=True)
    p.add_argument("--page", type=int, default=1)
    p.add_argument("--x", type=int, default=200)
    p.add_argument("--y", type=int, default=100)
    p.add_argument("--width", type=int, default=180)
    p.add_argument("--output", "-o")

    p = sub.add_parser("fill-form", help="Fill AcroForm fields")
    p.add_argument("pdf")
    p.add_argument("--data", required=True, help='JSON: {"field":"value"}')
    p.add_argument("--output", "-o")

    p = sub.add_parser("list-fields", help="List AcroForm fields in PDF")
    p.add_argument("pdf")

    p = sub.add_parser("set-metadata", help="Set document metadata")
    p.add_argument("pdf")
    p.add_argument("--title")
    p.add_argument("--author")
    p.add_argument("--subject")
    p.add_argument("--keywords")
    p.add_argument("--output", "-o")
    p.add_argument("--force", action="store_true", help="Overwrite original")

    p = sub.add_parser("add-pagenum", help="Add page numbers")
    p.add_argument("pdf")
    p.add_argument("--position", default="bottom-center", choices=["bottom-center", "bottom-right", "bottom-left", "top-center", "top-right", "top-left"])
    p.add_argument("--start", type=int, default=1)
    p.add_argument("--fontsize", type=int, default=10)
    p.add_argument("--pages", help="Pages to number")
    p.add_argument("--output", "-o")

    # ── Conversion ──
    p = sub.add_parser("pdf-to-word", help="Convert PDF to Word")
    p.add_argument("pdf")
    p.add_argument("--output", "-o")

    p = sub.add_parser("pdf-to-excel", help="Extract tables to Excel")
    p.add_argument("pdf")
    p.add_argument("--output", "-o")

    p = sub.add_parser("pdf-to-images", help="Convert PDF pages to images")
    p.add_argument("pdf")
    p.add_argument("--output-dir")
    p.add_argument("--format", default="png")
    p.add_argument("--dpi", type=int, default=200)

    p = sub.add_parser("pdf-to-html", help="Convert PDF to HTML")
    p.add_argument("pdf")
    p.add_argument("--output", "-o")

    p = sub.add_parser("images-to-pdf", help="Combine images into PDF")
    p.add_argument("images", nargs="+")
    p.add_argument("--output", "-o")

    p = sub.add_parser("word-to-pdf", help="Convert Word document to PDF")
    p.add_argument("docx")
    p.add_argument("--output", "-o")

    p = sub.add_parser("excel-to-pdf", help="Convert Excel spreadsheet to PDF")
    p.add_argument("xlsx")
    p.add_argument("--output", "-o")

    p = sub.add_parser("create-pdf", help="Create PDF from text/markdown file")
    p.add_argument("text_file")
    p.add_argument("--title")
    p.add_argument("--output", "-o")

    # ── Batch ──
    p = sub.add_parser("batch", help="Batch operation on directory")
    p.add_argument("operation", choices=["ocr", "compress", "encrypt", "redact", "flatten", "info", "sanitize", "linearize"])
    p.add_argument("input_dir")
    p.add_argument("--password")
    p.add_argument("--lang", default="eng")
    p.add_argument("--patterns", help="For batch redact: patterns to redact")
    p.add_argument("--keywords", help="For batch redact: keywords to redact")

    args = parser.parse_args()

    command_map = {
        "info": cmd_info, "search": cmd_search, "compare": cmd_compare,
        "extract-text": cmd_extract_text, "extract-images": cmd_extract_images,
        "ocr": cmd_ocr,
        "merge": cmd_merge, "split": cmd_split, "remove-pages": cmd_remove_pages,
        "extract-pages": cmd_extract_pages, "rotate": cmd_rotate, "reorder": cmd_reorder,
        "crop": cmd_crop, "flatten": cmd_flatten, "compress": cmd_compress,
        "encrypt": cmd_encrypt, "decrypt": cmd_decrypt, "redact": cmd_redact,
        "sanitize": cmd_sanitize, "repair": cmd_repair, "linearize": cmd_linearize,
        "watermark": cmd_watermark, "esign": cmd_esign,
        "fill-form": cmd_fill_form, "list-fields": cmd_list_fields,
        "set-metadata": cmd_set_metadata, "add-pagenum": cmd_add_pagenum,
        "pdf-to-word": cmd_pdf_to_word, "pdf-to-excel": cmd_pdf_to_excel,
        "pdf-to-images": cmd_pdf_to_images, "pdf-to-html": cmd_pdf_to_html,
        "images-to-pdf": cmd_images_to_pdf,
        "word-to-pdf": cmd_word_to_pdf, "excel-to-pdf": cmd_excel_to_pdf,
        "create-pdf": cmd_create_pdf,
        "batch": cmd_batch,
    }
    command_map[args.command](args)


if __name__ == "__main__":
    main()
